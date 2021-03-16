import PySimpleGUI as sg
import datetime
import json

import sys
from docx import Document
from os import path


class FillNewCopy:
    """
        A Class that based in a configuration file and a template file, will create a new Word document (docx).
        by replacing placeholders present in the configuration

        The configuration file should be a valid json file similar to:
        {
            "app_name": "Generate",
            "template_file": "file_template.docx",
            "file_prefix": "",
            "file_posfix": "",
            "date_format": "%d-%m-%Y",
            "fields": {
                "int_num": {
                    "label": "Internal #",
                    "type": "str",
                    "default": "",
                    "required": 1
                },
                ...
        }
    """

    APP_NAME = "FillNewCopy"
    DEFAULT_DATE_FORMAT = "%d-%m-%Y"

    def __init__(self, config_path: str = "config.json"):
        """
            It will start the app by loading configuration

            Parameters:
                config_path (str, optional): Configuration file path. Defaults to "config.json".
        """
        # Change here if you want to relocate you config file
        self.config = {}
        self.load_configuration(config_path)
        self.app_name = self.config.get('app_name', self.APP_NAME)

    def load_configuration(self, config_path: str):
        """
            Load configuration

            Parameters:
                config_path (str): File location for config file

            Returns:                int:Returning value
        """
        # Try to open configuration file
        try:
            with open(config_path) as json_file:
                config_data = json.load(json_file)
                self.config = config_data

            # Validate config data
            if "app_name" not in self.config or self.config.get("app_name") == "":
                self.config.update({"app_name": self.APP_NAME})

            if "template_file" not in self.config:
                sg.Popup("Opps!", "No template file provided in config file. Please make sure that 'template_file'"
                                  " exists in the configuration and provides a correct path.")
                sys.exit()

            if not path.exists(self.config.get("template_file")):
                sg.Popup("Opps!", f"No template file found.")
                sys.exit()

            if "date_format" not in self.config:
                self.config.update({"date_format": self.DEFAULT_DATE_FORMAT})

            if "fields" not in self.config:
                sg.Popup("Opps!", f"No fields configured in the config, please check documentation.")
                sys.exit()

        except FileNotFoundError:
            sg.Popup("Opps!", "Config file not found.")
            sys.exit()

        except json.decoder.JSONDecodeError:
            sg.Popup("Opps!", "Bad Json in config file!")
            sys.exit()

    def start(self):
        """
            Create and launch window
        """
        window_layout = self.build_gui()
        window = sg.Window(self.app_name, window_layout)
        fields = self.config.get("fields")
        while True:
            event, values = window.read()
            if event == sg.WIN_CLOSED or event == 'Cancel':  # if user closes window or clicks cancel
                break

            if event == 'Build':
                # Validate fields
                errors = False
                for (key, value) in values.items():
                    if key in fields:
                        errmsg = ""
                        if fields.get(key).get("type") == "str":
                            errmsg = self.validate_text_field(
                                fields.get(key), value)
                        if fields.get(key).get("type") == "int":
                            errmsg = self.validate_int_field(
                                fields.get(key), value)
                        if fields.get(key).get("type") == "list":
                            errmsg = self.validate_list_field(
                                fields.get(key), value)
                        if fields.get(key).get("type") == "date":
                            errmsg = self.validate_date_field(
                                fields.get(key), value)
                        if fields.get(key).get("type") == "textarea":
                            errmsg = self.validate_textarea_field(
                                fields.get(key), value)

                        if errmsg != "":
                            sg.Popup("Opps!", f"{errmsg}")
                            errors = True
                            break

                # Build document
                if not errors:
                    self.sanitize_values(values)
                    try:
                        filename = self.build_document(values)
                        sg.Popup(
                            "Congrats!", f"Your file ({filename}) was generated!")
                        break
                    except Exception:
                        e = sys.exc_info()[0]
                        sg.Popup(f"Problem generating your file. (Error: {e})")

    @staticmethod
    def sanitize_values(values: dict):
        """
            Due to the way lists works in PySimpleGUI,
            we need to make sure we parse them correctly

            Parameters:
                values (dict): Dictionary containing the values

        """
        for (key, value) in values.items():
            if isinstance(value, list):
                values.update({key: value[0]})

    def build_gui(self):
        """
            Build GUI based in fields loaded from
            the config

            NOTE: At this moment int works equally to
            str. Hopefully to be improved in the future

            Returns:
                List of elements to be used in the layout
        """
        # Build header
        layout = [[sg.Text(f"Welcome to {self.app_name}")], [sg.Text('')]]

        # Build form
        for (field_name, field) in (self.config.get("fields")).items():
            # By default we will use str as type
            if "type" not in field:
                field.update({"type": "str"})

            # Make sure we have a default value
            if "default" not in field:
                field.update({"default": ""})

            if field.get("type") == "str" or field.get("type") == "int":
                layout.append(self.build_string_field(field_name, field))
            elif field.get("type") == "date":
                layout.append(self.build_date_field(field_name, field))
            elif field.get("type") == "list":
                layout.append(self.build_list_field(field_name, field))
            elif field.get("type") == "textarea":
                layout.append(self.build_textarea_field(field_name, field))
            else:  # If not identified, just treat it as a str
                layout.append(self.build_string_field(field_name, field))

        # Build footer
        layout.append([sg.Text('')])
        layout.append([sg.Text('* Mandatory fields', text_color="Red")])
        layout.append([sg.Button('Build'), sg.Button('Cancel')])
        layout.append([sg.Text('')])
        return layout

    def build_string_field(self, field_name: str, field: dict):
        """
            Build list with elements of a string field

            Parameters:
                field_name (str): The field name
                field (dict): The field configuration

            Returns:
                list: List with the elements composing a text field
        """

        field_layout = [sg.Text(self.build_label_text(field_name, field), size=(15, 1)),
                        sg.InputText(field.get("default"), key=field_name)]

        return field_layout

    @staticmethod
    def validate_text_field(field: dict, value: str):
        """
            Validates text fields

        Args:
            field (dict): The field configuration
            value (str): input value

        Returns:
            str: empty string if no errors, otherwise error message
        """
        if field.get("required") and value.strip() == "":
            return f"{field.get('label')} is required!"
        return ""

    @staticmethod
    def validate_int_field(field: dict, value: str):
        """
            Validates int fields

        Args:
            field (dict): The field configuration
            value (str): input value

        Returns:
            str: empty string if no errors, otherwise error message
        """
        if field.get("required") and value.strip() == "":
            return f"{field.get('label')} is required!"

        try:
            temp = int(value)
        except ValueError:
            return f"{field.get('label')} should be a Number"

        return ""

    def build_date_field(self, field_name: str, field: dict):
        """
            Build list with elements of a date field

        Args:
            field_name (str): The field name
            field (dict): The field configuration

        Returns:
            list: List with the elements composing a date field
        """
        now = (datetime.datetime.now()).strftime(
            self.config.get("date_format"))
        field_layout = [sg.Text(self.build_label_text(field_name, field), size=(15, 1)),
                        sg.InputText(now, key=field_name,
                                     enable_events=False, visible=True),
                        sg.CalendarButton('Calendar', target=field_name,
                                          key='CALENDAR', format=(self.config.get("date_format")))]

        return field_layout

    def validate_date_field(self, field: dict, value: str):
        """
            Validates a date field

        Args:
            field (dict): The field configuration
            value (str): input value

        Returns:
            str: empty string if no errors, otherwise error message
        """
        if field.get("required") and value.strip() == "":
            return f"{field.get('label')} is required!"

        try:
            datetime.datetime.strptime(value, self.config.get("date_format"))
        except ValueError:
            return f"{field.get('label')} should be a date with the format provided in " \
                   f"config {self.config.get('date_format')}"

        return ""

    def build_list_field(self, field_name: str, field: dict):
        """
            Build list with elements of a list field

        Args:
            field_name (str): The field name
            field (dict): The field configuration

        Returns:
            list: List with the elements composing a list field
        """
        field_layout = [sg.Text(self.build_label_text(field_name, field), size=(15, 1)),
                        sg.Listbox(field.get("options"), default_values=field.get("default"), size=(20, 4),
                                   enable_events=False, key=field_name)]

        return field_layout

    @staticmethod
    def validate_list_field(field: dict, value: list):
        """
            Validates a list field

        Args:
            field (dict): The field configuration
            value (str): input value

        Returns:
            str: empty string if no errors, otherwise error message
        """

        if field.get("required") and len(value) == 0:
            return f"{field.get('label')} is required!"

        return ""

    def build_textarea_field(self, field_name: str, field: dict):
        field_layout = [sg.Text(self.build_label_text(field_name, field), size=(15, 1)),
                        sg.Multiline(field.get("default"), size=(30, 5), key=field_name)]

        return field_layout

    @staticmethod
    def validate_textarea_field(field: dict, value: str):
        """
            Validates a textarea field

        Args:
            field (dict): The field configuration
            value (str): input value

        Returns:
            str: empty string if no errors, otherwise error message
        """

        if field.get("required") and value.strip() == "":
            return f"{field.get('label')} is required!"

        return ""

    @staticmethod
    def build_label_text(field_name: str, field: dict):
        """
            Returns the label text

        Args:
            field_name (str): The field name
            field (dict): The field configuration

        Returns:
            str: formatted string to be used in the label
        """

        label = ""
        if "required" in field:
            label = " * " if field.get("required") else ""

        # If we don't have a label defined, used the field name
        if "label" not in field:
            field.update({"label": field_name.upper()})

        label += field["label"]

        return label

    def build_document(self, values: dict):
        """
            Tries to parse the document and substitute
        place holders by the input values

        Args:
            values (dict): Dictionary containing the 
            input values

        Returns:
            str: filename

        """
        doc = Document(self.config.get('template_file'))
        for section in doc.sections:
            # First Header
            header = section.header
            for p in header.paragraphs:
                for key, field in self.config.get('fields').items():
                    # Format key
                    formatted_key = f"<{key.upper()}>"
                    if formatted_key in p.text:
                        inline = p.runs
                        # Loop added to work with runs (strings with same style)
                        for i in range(len(inline)):
                            if formatted_key in inline[i].text:
                                text = inline[i].text.replace(
                                    formatted_key, values[key])
                                inline[i].text = text
            # Second Footer
            footer = section.footer
            for p in footer.paragraphs:
                for key, field in self.config.get("fields").items():
                    # Format key
                    formatted_key = f"<{key.upper()}>"
                    if formatted_key in p.text:
                        inline = p.runs
                        # Loop added to work with runs (strings with same style)
                        for i in range(len(inline)):
                            if formatted_key in inline[i].text:
                                text = inline[i].text.replace(
                                    formatted_key, values[key])
                                inline[i].text = text

        # Go by tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, field in self.config.get("fields").items():
                            # Format key
                            formatted_key = f"<{key.upper()}>"
                            if formatted_key in p.text:
                                inline = p.runs
                                # Loop added to work with runs (strings with same style)
                                for i in range(len(inline)):
                                    if formatted_key in inline[i].text:
                                        text = inline[i].text.replace(formatted_key, values[key])
                                        inline[i].text = text

        # Go by the rest of the document
        for p in doc.paragraphs:
            for key, field in self.config.get("fields").items():
                # Format key
                formatted_key = f"<{key.upper()}>"
                if formatted_key in p.text:
                    inline = p.runs
                    # Loop added to work with runs (strings with same style)
                    for i in range(len(inline)):
                        if formatted_key in inline[i].text:
                            text = inline[i].text.replace(formatted_key, values[key])
                            inline[i].text = text

        # By default filename will be the template filename with copy_ before
        filename = f"copy_{self.config.get('template_file')}"
        if "filename" in self.config:
            if "type" in self.config.get("filename"):
                # We can have 2 types static value or based in a field
                if self.config.get("filename").get("type") == "static" and "value" in self.config.get("filename"):
                    filename = self.config.get("filename").get("value")
                elif self.config["filename"]["type"] == "field" and "value" in self.config["filename"]:
                    filename = values.get(self.config.get("filename").get("value"))

        # Make sure we have a prefix populated
        if "file_prefix" not in self.config:
            self.config.update({"file_prefix": ""})

        # Make sure we have a posfix populated
        if "file_posfix" not in self.config:
            self.config.update({"file_posfix": ""})

        filename = self.config.get("file_prefix") + filename + self.config.get("file_posfix")

        doc.save(f"{filename}.docx")

        return f"{filename}.docx"


if __name__ == "__main__":
    rn = FillNewCopy()
    rn.start()

